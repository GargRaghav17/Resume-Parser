from flask import Flask, request, send_file, abort, render_template_string
import os
import pandas as pd
import pandas as pd
from pyresparser import ResumeParser
from typing import List
import pandas as pd
import os
import warnings
from docx import Document
import PyPDF2
from doc2docx import convert
import subprocess

warnings.filterwarnings('ignore')

app = Flask(__name__)

def convert_doc_to_docx(doc_path):
    # Construct the path to the .docx file
    docx_path = os.path.splitext(doc_path)[0] + '.docx'
    
    # Convert the .doc file to .docx
    subprocess.run(['soffice', '--headless', '--convert-to', 'docx', doc_path], check=True)
    
    return docx_path

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        try:
            # Create the 'uploads' directory if it doesn't exist
            uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
            os.makedirs(uploads_dir, exist_ok=True)

            # Initialize an empty DataFrame to hold the data from all PDFs
            df = pd.DataFrame()

            # Get the PDF files from the user
            pdf_files = request.files.getlist('pdf_file')

            for pdf_file in pdf_files:
                # Save the PDF file to the server
                pdf_file_path = os.path.join(uploads_dir, pdf_file.filename)
                pdf_file.save(pdf_file_path)

                # Process the PDF file and append the data to df
                df = pd.concat([df, pd.DataFrame([process_pdf(pdf_file_path, uploads_dir, df)])], ignore_index=True)

            # Save the DataFrame to an Excel file
            excel_file_path = os.path.join(uploads_dir, 'output.xlsx')
            df.to_excel(excel_file_path, index=False)

            # Return the Excel file to the user
            return send_file(excel_file_path, as_attachment=True)
        except Exception as e:
            # Log the error and return a meaningful error message
            app.logger.error(f"Error processing PDF files: {e}")
            abort(500, "An error occurred while processing the PDF files.")


    return render_template_string("""
    <!doctype html>
    <html lang="en">
      <head>
        <!-- Required meta tags -->
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1">

        <!-- Bootstrap CSS -->
        <link href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css" rel="stylesheet">

        <title>Resume Parser</title>
      </head>
      <body>
        <div class="container">
          <h1 class="mt-5">Upload Your File(s)</h1>
          <p class="lead">Please ensure your files are in (.pdf/ .docx/ .txt) format for best results.</p>
          <form method="post" enctype="multipart/form-data" class="mt-4">
            <div class="custom-file">
              <input type="file" class="custom-file-input" id="pdfFile" name="pdf_file" multiple>
              <label class="custom-file-label" for="pdfFile">Choose file</label>
            </div>
            <button type="submit" class="btn btn-primary mt-3">Upload</button>
          </form>
        </div>

        <!-- Optional JavaScript -->
        <!-- jQuery first, then Popper.js, then Bootstrap JS -->
        <script src="https://code.jquery.com/jquery-3.5.1.slim.min.js"></script>
        <script src="https://cdn.jsdelivr.net/npm/bootstrap@4.5.2/dist/js/bootstrap.bundle.min.js"></script>

        <script>
        // Update the label of the file input with the names of the selected files
        document.getElementById('pdfFile').addEventListener('change', function() {
            var label = Array.from(this.files).map(file => file.name).join(', ');
            if (this.files.length > 1) {
                label = this.files.length + ' files selected';
            }
            this.nextElementSibling.textContent = label;
        });
        </script>
      </body>
    </html>
    """)

def process_pdf(pdf_path, uploads_dir, df):
    df = pd.DataFrame()
    filed = pdf_path

    cv_file = filed

    if cv_file.endswith('.docx'):
        # filed = os.path.join(folder_path, cv_file)

        try:
            doc = Document()
            with open(filed, 'r') as file:
                text = file.read()
                cleaned_text = ''.join(str(ch) if not isinstance(ch, int) and ch.isprintable() else '' for ch in text)
                doc.add_paragraph(cleaned_text)

            doc.save("text.docx")
            data = ResumeParser("text.docx").get_extracted_data()
            # df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
            return data

        except:
            data = ResumeParser(filed).get_extracted_data()
            # df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
            return data


    elif cv_file.endswith('.doc'):
        # filed = os.path.join(folder_path, cv_file)
        convert(filed)
        filedd = filed + 'x'
        os.remove(filed)
        # docx_file = os.path.splitext(filed)[0] + '.docx'
        # subprocess.run(['lowriter', '--convert-to', 'docx', filed], check=True)
        # # subprocess.call(['soffice', '--headless', '--convert-to', 'docx', filename])
        # # filedd = filed + 'x'
        # filed = docx_file
        # print(f"File path: {filed}")

        # Works on Linux
        # try:
        #     subprocess.run(['soffice', '--headless', '--convert-to', 'docx', filed])
        #     # print(f"File path after conversion: {filed}")

        # except:
        #     # print(f"File path after error: {filed}")
        #     None
        # Till here

        try:
            doc = Document()
            with open(filedd, 'r') as file:
                text = file.read()
                cleaned_text = ''.join(str(ch) if not isinstance(ch, int) and ch.isprintable() else '' for ch in text)
                doc.add_paragraph(cleaned_text)

            doc.save("text.docx")
            data = ResumeParser("text.docx").get_extracted_data()
            # df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
            return data

        except:
            data = ResumeParser(filedd).get_extracted_data()
            # df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
            return data


    elif cv_file.endswith('.txt'):
        # filed = os.path.join(folder_path, cv_file)

        try:
            doc = Document()
            with open(filed, 'r', encoding='utf-8') as file:
                text = file.read()
                cleaned_text = ''.join(ch for ch in text if ch.isprintable())
                doc.add_paragraph(cleaned_text)

            doc.save("text.docx")
            data = ResumeParser("text.docx").get_extracted_data()
            # df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
            return data
            
        except Exception as e:
            data = ResumeParser(filed).get_extracted_data()
            # df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
            return data


    elif cv_file.endswith('.pdf'):
        # filed = os.path.join(folder_path, cv_file)

        try:
            doc = Document()
            with open(filed, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ' '.join([reader.pages[i].extract_text() for i in range(len(reader.pages))])
                cleaned_text = ''.join(ch for ch in text if ch.isprintable())
                doc.add_paragraph(cleaned_text)

            doc.save("text.docx")
            data = ResumeParser("text.docx").get_extracted_data()
            # df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
            return data
            
        except Exception as e:
            data = ResumeParser(filed).get_extracted_data()
            # df = pd.concat([df, pd.DataFrame([data])], ignore_index=True)
            return data
    
    # output_file_path = os.path.join(uploads_dir, 'output.xlsx')
    # df.to_excel(output_file_path, index=False)

if __name__ == '__main__':
    app.run()